from pathlib import Path
import pickle
from typing import List, Dict, Optional
import pdfplumber
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from striprtf.striprtf import rtf_to_text
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer, util
from paddleocr import PaddleOCR
ocr = PaddleOCR(lang='ru')
POPPLER_PATH = r"C:\Program Files\poppler\Library\bin"
RAW_DIR = Path("data/raw")
EXTRACTED_DIR = Path("data/extracted")
VECTOR_STORE_PATH = Path("data/vector_store")
RAW_DIR.mkdir(parents=True, exist_ok=True)
EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)
VECTOR_STORE_PATH.mkdir(parents=True, exist_ok=True)
EMBEDDING_MODEL_NAME = "models/all-mpnet-base-v2"
embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
CHUNK_SIZE = 800
def extract_pdf(path: Path) -> Optional[str]:
    try:
        with pdfplumber.open(path) as pdf:
            pieces = [page.extract_text() for page in pdf.pages if page.extract_text()]
        return "\n\n".join(pieces) or None
    except:
        return None
def extract_pdf_ocr(path: Path) -> Optional[str]:
    try:
        images = convert_from_path(str(path), poppler_path=POPPLER_PATH)
        collected = []
        for img in images:
            result = ocr.ocr(img)
            for line in result:
                for box in line:
                    text = box[1][0]
                    if text.strip():
                        collected.append(text)
        return "\n".join(collected) or None
    except Exception as e:
        return None

def extract_docx(path: Path) -> Optional[str]:
    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append("[TABLE] " + row_text)
        return "\n\n".join(parts) or None
    except:
        return None

def extract_rtf(path: Path) -> Optional[str]:
    try:
        content = path.read_text(encoding="utf-8", errors="ignore")
        text = rtf_to_text(content)
        return text.strip() or None
    except:
        return None

def extract_text(path: Path) -> Optional[str]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        t = extract_pdf(path)
        if not t:
            t = extract_pdf_ocr(path)
        return t
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".rtf":
        return extract_rtf(path)
    return None

# ===================== Chunking =====================

def recursive_split(text: str, size: int) -> List[str]:
    sentences = util.split_into_sentences(text)
    chunks = []
    buf = ""
    for s in sentences:
        if len(buf) + len(s) + 1 <= size:
            buf = (buf + " " + s).strip()
        else:
            if buf:
                chunks.append(buf)
            buf = s
    if buf:
        chunks.append(buf)
    return chunks

# ===================== Основной пайплайн =====================

def process_all():
    file_paths = list(RAW_DIR.rglob("*.pdf")) + list(RAW_DIR.rglob("*.docx")) + list(RAW_DIR.rglob("*.rtf"))
    all_chunks = []
    all_payloads = []
    counter = 1

    for path in file_paths:
        text = extract_text(path)
        if not text:
            continue

        # Сохраняем сырой текст
        out_file = EXTRACTED_DIR / f"{path.stem}.txt"
        out_file.write_text(text, encoding="utf-8", errors="ignore")

        chunks = recursive_split(text, CHUNK_SIZE)
        for i, ch in enumerate(chunks):
            all_chunks.append(ch)
            all_payloads.append({
                "id": counter,
                "chunk": i,
                "source": path.name,
                "text": ch
            })
            counter += 1

    print(f"Total chunks: {len(all_chunks)}")
    return all_chunks, all_payloads

# ===================== Векторизация и FAISS =====================

def build_faiss_index(chunks: List[str], payloads: List[Dict]):
    embs = embedding_model.encode(chunks, convert_to_numpy=True)
    # Нормализация для cosine similarity
    embs = embs / np.linalg.norm(embs, axis=1, keepdims=True)

    dim = embs.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embs)

    faiss.write_index(index, str(VECTOR_STORE_PATH / "faiss.index"))
    np.save(VECTOR_STORE_PATH / "embeddings.npy", embs)

    with open(VECTOR_STORE_PATH / "payloads.pkl", "wb") as f:
        pickle.dump(payloads, f)

    print("FAISS index built and saved.")
# ===================== RUN =====================

if __name__ == "__main__":
    chunks, payloads = process_all()
    build_faiss_index(chunks, payloads)
    print("Pipeline completed.")
